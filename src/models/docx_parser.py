from typing import Optional, Union
import os
import re
import docx

# Arabic paragraph letters used in Jordanian law (in order: أ ب ج د هـ و ز ح ط ي)
# Plus continuation letters sometimes seen: ك ل م ن س ع ف ص ق
# We explicitly list them all instead of using a range [أ-ي] which is unreliable in Unicode
_AR_LETTER = (
    r"(?:هـ|[أبتثجحخدذرزسشصضطظعغفقكلمنهوي])"
)

# Matches: "أ- text"  or  "هـ- text"  (standalone paragraph)
# The (?!\d+\s*-) negative lookahead prevents matching "أ- 1- ..." as a paragraph
_PARA_RE = re.compile(
    r"^\u200b?\u200c?\u200d?"           # optional zero-width chars (ZWNJ, ZWJ, BOM)
    r"(" + _AR_LETTER + r")"            # capture letter (including هـ)
    r"\s*-\s*"                           # dash, optional spaces
    r"(?!\d+\s*-)"                       # NOT followed by digit-dash (that's a combined line)
    r"(.*)",                             # rest of text
    re.DOTALL
)

# Matches: "أ- 1- text"  or  "أ-1- text"  or  "أ-1 - text" (combined letter+item)
_COMBINED_RE = re.compile(
    r"^\u200b?\u200c?\u200d?"           # optional zero-width chars
    r"(" + _AR_LETTER + r")"            # capture letter
    r"\s*-\s*"                           # dash
    r"(\d+)"                             # capture item number
    r"\s*-\s*"                           # dash (with optional spaces on either side)
    r"(.*)",                             # rest of text
    re.DOTALL
)

# Matches: "1- text"  or  "15- text"
_ITEM_RE = re.compile(r"^(\d+)-\s*(.*)", re.DOTALL)

# Article header: "المادة 5   01-01-2015"
# The date part uses \xa0 (non-breaking space) as separator in the actual file
_ARTICLE_RE = re.compile(
    r"^المادة\s+(\d+)"                  # article number (Western digits)
    r"(?:[\s\xa0]+(\d{2,4}-\d{1,2}-\d{1,4}))?"  # optional effective date
)


class DocxParser:
    @staticmethod
    def parse_law(file_path):
        """
        Parses a Law docx file.
        Returns a dictionary representing the law structure:
        {
            "title": str,
            "law_number": int,
            "law_year": int,
            "articles": [
                {
                    "number": int,
                    "effective_date": str,
                    "text": str,
                    "paragraphs": [
                        {
                            "letter": str,
                            "text": str,
                            "items": [
                                {
                                    "number": int,
                                    "text": str
                                }
                            ]
                        }
                    ]
                }
            ]
        }
        """
        doc = docx.Document(file_path)
        filename = os.path.basename(file_path)

        # ── Extract law metadata from header paragraphs ──────────────────────
        title      = "قانون ضريبة الدخل"
        law_number = 34
        law_year   = 2014

        # Try to read from filename first
        fn_match = re.search(r"رقم\s+(\d+)\s+لسنة\s+(\d+)", filename)
        if fn_match:
            law_number = int(fn_match.group(1))
            law_year   = int(fn_match.group(2))

        # Scan first 20 non-empty paragraphs for metadata
        non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]
        for p in non_empty_paras[:20]:
            t = p.text.strip()
            if t.startswith("الرقم:"):
                m = re.search(r"\d+", t)
                if m:
                    law_number = int(m.group())
            elif t.startswith("السنة:"):
                m = re.search(r"\d+", t)
                if m:
                    law_year = int(m.group())
            elif "قانون" in t and "لسنة" in t and not title.endswith(t):
                # e.g. "قانون رقم 34 لسنة 2014 (قانون ضريبة الدخل لسنة 2014) وتعديلاته"
                title = t.split("(")[-1].split(")")[0] if "(" in t else t

        # ── Walk through paragraphs and build article list ───────────────────
        articles: list[dict] = []
        current_article: Optional[dict] = None
        current_paragraph: Optional[dict] = None

        # Helper to iterate through paragraphs and tables in document order
        from docx.text.paragraph import Paragraph
        from docx.table import Table

        def iter_block_items(parent):
            if isinstance(parent, docx.document.Document):
                parent_elm = parent.element.body
            else:
                parent_elm = parent._tc
            for child in parent_elm.iterchildren():
                if child.tag.endswith('p'):
                    yield Paragraph(child, parent)
                elif child.tag.endswith('tbl'):
                    yield Table(child, parent)

        for block in iter_block_items(doc):
            if isinstance(block, Table):
                # If we are currently parsing Article 2, parse definitions from the table
                if current_article and current_article["number"] == 2:
                    for row in block.rows:
                        seen_cells = []
                        for cell in row.cells:
                            c_text = cell.text.strip()
                            if c_text and c_text not in seen_cells:
                                seen_cells.append(c_text)
                        
                        if len(seen_cells) >= 2:
                            term = seen_cells[0]
                            if seen_cells[1] == ":" and len(seen_cells) >= 3:
                                definition = " ".join(seen_cells[2:])
                            else:
                                definition = " ".join(seen_cells[1:])
                            
                            definition = definition.strip().lstrip(":").strip()
                            
                            definition_line = f"{term}: {definition}"
                            if current_article["text"]:
                                current_article["text"] += "\n" + definition_line
                            else:
                                current_article["text"] = definition_line
                                
                            if not current_article["paragraphs"]:
                                current_paragraph = {
                                    "letter": "عام",
                                    "text":   definition_line,
                                    "items":  []
                                }
                                current_article["paragraphs"].append(current_paragraph)
                            else:
                                if current_paragraph:
                                    current_paragraph["text"] += "\n" + definition_line
                continue

            # It's a paragraph block
            text = block.text.strip()
            # Remove common zero-width characters at start of line that can confuse matching
            text_clean = text.lstrip("\u200b\u200c\u200d\ufeff")
            if not text_clean:
                continue

            # ── 1. Article header ─────────────────────────────────────────
            art_match = _ARTICLE_RE.match(text_clean)
            if art_match:
                art_num  = int(art_match.group(1))
                eff_date = art_match.group(2) if art_match.group(2) else "2015-01-01"

                current_article = {
                    "number":         art_num,
                    "effective_date": eff_date,
                    "text":           "",
                    "paragraphs":     []
                }
                articles.append(current_article)
                current_paragraph = None
                continue

            if current_article is None:
                continue  # still in preamble before first article

            # Append raw text to article
            if current_article["text"]:
                current_article["text"] += "\n" + text_clean
            else:
                current_article["text"] = text_clean

            # ── 2. Combined line: "أ- 1- text"  or  "أ-1 - text" ─────────
            combined_match = _COMBINED_RE.match(text_clean)
            if combined_match:
                letter    = combined_match.group(1)
                item_num  = int(combined_match.group(2))
                item_text = combined_match.group(3).strip()

                # Check if we already have a paragraph with this letter
                existing = next(
                    (pg for pg in current_article["paragraphs"] if pg["letter"] == letter),
                    None
                )
                if existing:
                    # Add the item to the existing paragraph
                    existing["items"].append({"number": item_num, "text": item_text})
                    current_paragraph = existing
                else:
                    # New paragraph+item
                    current_paragraph = {
                        "letter": letter,
                        "text":   item_text,
                        "items":  [{"number": item_num, "text": item_text}]
                    }
                    current_article["paragraphs"].append(current_paragraph)
                continue

            # ── 3. Standalone paragraph letter: "أ- text" ────────────────
            para_match = _PARA_RE.match(text_clean)
            if para_match:
                letter    = para_match.group(1)
                para_text = para_match.group(2).strip()

                current_paragraph = {
                    "letter": letter,
                    "text":   para_text,
                    "items":  []
                }
                current_article["paragraphs"].append(current_paragraph)
                continue

            # ── 4. Standalone item: "1- text" ─────────────────────────────
            item_match = _ITEM_RE.match(text_clean)
            if item_match and current_paragraph is not None:
                item_num  = int(item_match.group(1))
                item_text = item_match.group(2).strip()
                current_paragraph["items"].append({"number": item_num, "text": item_text})
                continue

            # ── 5. Plain continuation text ─────────────────────────────────
            if not current_article["paragraphs"]:
                # Article with no letter-paragraphs yet — create a generic one
                default_para = {
                    "letter": "عام",
                    "text":   text_clean,
                    "items":  []
                }
                current_article["paragraphs"].append(default_para)
                current_paragraph = default_para
            else:
                # Append to current paragraph text (continuation line)
                if current_paragraph:
                    current_paragraph["text"] += "\n" + text_clean

        return {
            "title":      title,
            "law_number": law_number,
            "law_year":   law_year,
            "articles":   articles
        }

    @staticmethod
    def parse_judgment(file_path):
        """
        Parses a Judgment docx file.
        Returns a dictionary representing the judgment structure:
        {
            "ruling_id": str,
            "ruling_number": int,
            "ruling_year": int,
            "case_number": str,
            "court": str,
            "court_type": str,   # "administrative" | "tax_first" | "cassation" | "appeal_tax"
            "date": str,
            "outcome": str,
            "subject": str,
            "full_text": str,
            "citations": [
                {
                    "law_number": int,
                    "law_year": int,
                    "law_name": str,
                    "article_number": int,
                    "paragraph_letter": str | None,
                    "item_number": int | None,
                    "citation_text": str,   # النص الحرفي من الحكم
                    "resolved": bool        # True لو القانون موجود عندنا
                }
            ]
        }
        """
        doc = docx.Document(file_path)
        filename = os.path.basename(file_path)

        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        # ── Metadata: ruling number & year from filename ─────────────────────
        ruling_number = 0
        ruling_year   = 0
        fn_match = re.search(r"رقم\s+(\d+)\s+لسنة\s+(\d+)", filename)
        if fn_match:
            ruling_number = int(fn_match.group(1))
            ruling_year   = int(fn_match.group(2))

        ruling_id   = filename.replace(".docx", "")
        case_number = ""
        court       = "المحكمة الضريبية"
        court_type  = "tax_first"
        date        = ""
        outcome     = ""
        subject     = ""

        # ── Scan first 15 paragraphs for structured metadata ─────────────────
        for p in doc.paragraphs[:15]:
            t = p.text.strip()
            if not t:
                continue

            # Case number patterns
            for pat in [
                r"رقم (?:الدعوى|القضية)[:\s]*([\d\s/]+)",
                r"رقم القضية[:\s]*([\d\s/]+)",
                r"رقم الدعوى[:\s]*([\d\s/]+)",
            ]:
                m = re.search(pat, t)
                if m and not case_number:
                    case_number = m.group(1).strip().replace(" ", "")

            # Court detection
            if "المحكمة الإدارية العليا" in t:
                court      = "المحكمة الإدارية العليا"
                court_type = "administrative"
            elif "محكمة التمييز" in t:
                court      = "محكمة التمييز الأردنية"
                court_type = "cassation"
            elif "محكمة الاستئناف الضريبية" in t:
                court      = "محكمة الاستئناف الضريبية"
                court_type = "appeal_tax"
            elif "محكمة البداية الضريبية" in t or "البداية الضريبية" in t:
                court      = "محكمة البداية الضريبية"
                court_type = "tax_first"

            # Date
            for pat in [r"تاريخ الفصل:\s*([\d\-]+)", r"بتاريخ\s+([\d\s/]+(?:\d{4}))"]:
                m = re.search(pat, t)
                if m and not date:
                    raw = m.group(1).strip()
                    # Normalise "8 / 4 / 2024" → "2024-04-08"
                    parts = re.split(r"[\s/]+", raw)
                    if len(parts) == 3 and len(parts[-1]) == 4:
                        date = f"{parts[-1]}-{parts[1].zfill(2)}-{parts[0].zfill(2)}"
                    else:
                        date = raw

        # Outcome heuristic from last few paragraphs
        for p in doc.paragraphs[-10:]:
            t = p.text.strip()
            if "نقرر رد الطعن" in t or "رد الطعن" in t:
                outcome = "رد الطعن"
            elif "نقض القرار" in t:
                outcome = "نقض القرار"
            elif "إلغاء" in t and "قرار" in t:
                outcome = "إلغاء القرار"
            elif "قبول الدعوى" in t:
                outcome = "قبول الدعوى"

        # Fallbacks
        if not case_number:
            if ruling_number and ruling_year:
                case_number = f"{ruling_number}/{ruling_year}"
            else:
                case_number = ruling_id

        # ── Arabic letter set ─────────────────────────────────────────────────
        _AR = r"[أبتثجحخدذرزسشصضطظعغفقكلمنهوي]"

        # ── Known laws keyword map ────────────────────────────────────────────
        _LAW_KEYWORD_MAP = {
            "ضريبة الدخل":    (34,  2014, "قانون ضريبة الدخل"),
            "المبيعات":       (6,   1994, "قانون الضريبة العامة على المبيعات"),
            "أصول المحاكمات": (4,   1985, "قانون أصول المحاكمات المدنية"),
        }

        def _resolve_law(text_around: str) -> tuple:
            """
            Detect which law is being referenced from surrounding text.
            Returns (law_number, law_year, law_name, resolved).
            """
            # Explicit: "قانون ... رقم (34) لسنة 2014"
            m = re.search(
                r"قانون\s*([\w\s]+?)\s*رقم\s*[\(]?\s*(\d+)\s*[\)]?\s*لسنة\s*(\d{4})",
                text_around,
            )
            if m:
                name = f"قانون {m.group(1).strip()}"
                return int(m.group(2)), int(m.group(3)), name, True

            # Keyword match
            for kw, (num, yr, name) in _LAW_KEYWORD_MAP.items():
                if kw in text_around:
                    return num, yr, name, True

            # Default to income tax
            return 34, 2014, "قانون ضريبة الدخل", False

        # ── Citation collector ────────────────────────────────────────────────
        citations: list[dict] = []

        def _ctx(start: int, end: int, window: int = 200) -> str:
            """Surrounding context window for law resolution."""
            return full_text[max(0, start - window): end + window]

        def _add(art_num: int, para: "str|None", item: "int|None",
                 raw: str, ctx: str):
            law_num, law_yr, law_name, resolved = _resolve_law(ctx)
            citations.append({
                "law_number":       law_num,
                "law_year":         law_yr,
                "law_name":         law_name,
                "article_number":   art_num,
                "paragraph_letter": para,
                "item_number":      item,
                "citation_text":    raw.strip(),
                "resolved":         resolved,
            })

        # ── Pattern A: "المادة 11/و/2" — article / paragraph / item ─────────
        for m in re.finditer(
            r"المادة\s*[\(]?\s*(\d+)\s*/\s*(" + _AR + r")\s*/\s*(\d+)\s*[\)]?",
            full_text,
        ):
            _add(int(m.group(1)), m.group(2), int(m.group(3)),
                 m.group(0), _ctx(m.start(), m.end()))

        # ── Pattern B: "المادة (70/أ)" — article / paragraph ─────────────────
        for m in re.finditer(
            r"المادة\s*[\(]?\s*(\d+)\s*/\s*(" + _AR + r")\s*[\)]?",
            full_text,
        ):
            # Skip spans already covered by Pattern A
            span_text = m.group(0)
            already = any(
                c["article_number"] == int(m.group(1))
                and c["paragraph_letter"] == m.group(2)
                and c["item_number"] is not None
                and c["citation_text"].startswith(span_text[:15])
                for c in citations
            )
            if not already:
                _add(int(m.group(1)), m.group(2), None,
                     span_text, _ctx(m.start(), m.end()))

        # ── Pattern C: "المادتين (12) و (46)" ────────────────────────────────
        for m in re.finditer(
            r"المادت(?:ين|ان)\s*[\(]?\s*(\d+)\s*[\)]?\s*و\s*[\(]?\s*(\d+)\s*[\)]?",
            full_text,
        ):
            ctx = _ctx(m.start(), m.end())
            _add(int(m.group(1)), None, None, m.group(0), ctx)
            _add(int(m.group(2)), None, None, m.group(0), ctx)

        # ── Pattern D: "المواد (3) و (12) و (46)" ────────────────────────────
        for m in re.finditer(
            r"المواد\s*((?:[\(]?\s*\d+\s*[\)]?\s*(?:و|،|,)\s*)+[\(]?\s*\d+\s*[\)]?)",
            full_text,
        ):
            ctx = _ctx(m.start(), m.end())
            for n in re.findall(r"\d+", m.group(1)):
                _add(int(n), None, None, m.group(0), ctx)

        # ── Pattern E: Plain "المادة (N)" or "المادة N" ──────────────────────
        for m in re.finditer(r"المادة\s*[\(]?\s*(\d+)\s*[\)]?", full_text):
            art = int(m.group(1))
            # Skip if a more specific citation already covers this article in this span
            already = any(
                c["article_number"] == art
                and (c["paragraph_letter"] is not None or c["item_number"] is not None)
                and c["citation_text"].startswith(m.group(0)[:10])
                for c in citations
            )
            if not already:
                _add(art, None, None, m.group(0), _ctx(m.start(), m.end()))

        # ── De-duplicate: keep first occurrence of each (law, art, para, item) ─
        seen: set[tuple] = set()
        unique: list[dict] = []
        for c in citations:
            key = (c["law_number"], c["law_year"],
                   c["article_number"], c["paragraph_letter"], c["item_number"])
            if key not in seen:
                seen.add(key)
                unique.append(c)

        # Sort: article → paragraph letter → item number
        unique.sort(key=lambda c: (
            c["article_number"],
            c["paragraph_letter"] or "",
            c["item_number"] or 0,
        ))

        return {
            "ruling_id":     ruling_id,
            "ruling_number": ruling_number,
            "ruling_year":   ruling_year,
            "case_number":   case_number,
            "court":         court,
            "court_type":    court_type,
            "date":          date if date else "0000-00-00",
            "outcome":       outcome,
            "subject":       subject,
            "title":         ruling_id,
            "full_text":     full_text,
            "citations":     unique,
        }

    @staticmethod
    def _extract_law_number(filename: str) -> int:
        """Extract law number from filename, defaulting to 34."""
        m = re.search(r"رقم\s+(\d+)\s+لسنة", filename)
        return int(m.group(1)) if m else 34
